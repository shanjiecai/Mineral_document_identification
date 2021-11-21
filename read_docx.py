import re
import docx
from openpyxl import Workbook,load_workbook
from docx import Document
import requests
import json
import os
import jieba
# from win32com import client
from hanlp_ner2 import *
from string import digits
from get_keywords import *
import jpype
import pprint
from es import upload_data


# def doc2Docx(root, fileName):
#     fileName = os.path.join(root,fileName)
#     word = client.Dispatch("Word.Application")
#     doc = word.Documents.Open(fileName)
#     doc.SaveAs(fileName + "x", 16, False, "", True, "", False, False, False, False)
#     # os.remove(fileName)
#     doc.Close()
#     word.Quit()
#     return fileName+"x"


def get_mineral_name():
    wb = load_workbook("矿产名称.xlsx")
    sheets = wb.sheetnames
    sheet = wb[sheets[0]]
    mineral_list = []
    for index,row in enumerate(sheet.rows):
        if index==0:
            continue
        # print([cell.value for cell in row])
        mineral_list.extend(cell.value for cell in row)
    return mineral_list


def get_mineral_name_simple():
    f = open("矿产名称.txt",'r',encoding="utf-8")
    mineral_list = []
    for index, row in enumerate(f.readlines()):
        # print([cell.value for cell in row])
        mineral_list.append(row.strip())
    return mineral_list

def preprocess(text):
    return text


def ner(text):
    # 指定请求参数格式为json
    request_url = "https://nlpapi.aminer.cn/ner/"
    headers = {'Content-Type': 'application/json'}
    data = {
        "text": text
    }
    response = requests.post(request_url, headers=headers, data=json.dumps(data))
    if response:
        return response.json()['data']


# 获取中文字符
def get_chinese(text):
    pre = re.compile(u'[\u4e00-\u9fa5]')
    res = re.findall(pre, text)
    res1 = ''.join(res)
    return res1


# 文本清洗
def clean_text(text):
    # 需要自定义在医疗政策领域的词典
    wordlist = jieba.cut(text)
    # 去除停用词和长度小于2的词语
    wordlist = [w for w in wordlist if w not in stopwords and len(w) > 1]
    # 将中文数据组织成类似西方于洋那样，词语之间以空格间隔
    document =  "".join(wordlist)
    return document


def read_stopwords():
    # stopword_files = [
    #     'baidu_stopwords.txt', 'cn_stopwords.txt', 'hit_stopwords.txt',
    #     'my_stopwords.txt', 'scu_stopwords.txt', 'more_stopwords.txt'
    # ]
    stopword_files = ['stopword1.txt','stopword2.txt']
    stopwords_list = []
    for stopwords in stopword_files:
        f = open('./stopwords/' + stopwords, encoding='utf-8')
        for l in f.readlines():
            stopwords_list.append(l.strip())
    return stopwords_list


def word_segment(text):
    # 指定请求参数格式为json
    request_url = "https://nlpapi.aminer.cn/word_segmentation/word_segmentation"
    headers = {'Content-Type': 'application/json'}
    data = {
        "text": text,
        "cut_all": False,
        "HMM": True
    }
    response = requests.post(request_url, headers=headers, data=json.dumps(data))
    if response:
        # print(response.json()['data'])
        return response.json()['data']


from nltk.tokenize import RegexpTokenizer
def SplitSentence(content):
    tokenizer = RegexpTokenizer(".*?[。！？；，]") #就是以[]中的符号为标识分割的
    rst = tokenizer.tokenize(content)# list
    return rst


es_dataset = []
def read_docx(path):
    es_data = {}
    print(f"------------{path}------------")
    print("识别开始")
    document = Document(path)
    es_data['filename'] = path.split("/")[-1]
    para_list = []
    para_index = 1
    print("开始识别关键词")
    full_content = "\n".join(paragraph.text.strip() for paragraph in document.paragraphs)
    es_data["full_content"] = full_content
    keywords = get_keywords_aminer(full_content)
    print(keywords)
    keyword_list = []
    for k in keywords:
        keyword_list.append(k[0])
    es_data["keyword"] = ",".join(keyword_list)
    es_data["entity"] = {}
    for paragraph in document.paragraphs:
        mineral_tag = 0
        classify_tag = 0
        if paragraph.text.strip():
            # ner(paragraph.text.strip())
            word_segment_list = jieba.cut(paragraph.text.strip())
            for w in word_segment_list:
                if w in mineral_name:
                    mineral_tag=1
                if w in classify_name:
                    classify_tag=1
            if classify_tag and mineral_tag:
                print(f"第{para_index}段：", paragraph.text.strip())
                ner_dict = ner_dict_result(paragraph.text.strip())
                print("包含的实体为：\n", ner_dict)
                para_index += 1
                # sentence_list = SplitSentence(paragraph.text.strip())
                # for sent_index, sentence in enumerate(sentence_list):
                #     print(f"第{sent_index}句实体：", ner(sentence))
                para_list.append(paragraph.text.strip())
                es_data["entity"] = ner_dict
    es_dataset.append(es_data)
    return para_list


def recognize(text):
    # print(text)
    sentence_list = SplitSentence(text)

    result = {}
    mineral_tag = 0
    for sentence in sentence_list:
        word_segment_list = list(jieba.cut(sentence))
        # print(word_segment_list)
        for w in word_segment_list:
            if w in mineral_name:
                if w in result:
                    continue
                mineral_tag = 1
                mineral = w
                result[w] = {}
            if w in classify_name and mineral_tag:
                reg = re.compile(f"(?<={w}).*?\d+")
                match = reg.search(sentence)
                num = match.group(0)
                result[mineral][w] = num
    print("------该报告识别结果------")
    print(result)
    return result


def select_para(para_list):
    while True:
        para = input("请输入你认为准确的段落序号：(q键退出)\n")
        if para == "q":
            break
        print(para)
        recognize(para_list[int(para)-1])



if __name__ == "__main__":
    # 加在停用词
    stopwords = read_stopwords()
    # # 加在hanlp模型
    # ha_recognizer = train(PKU199801_TRAIN)
    mineral_name = get_mineral_name_simple()
    classify_name = ['推断', '控制', '探明', 'TM', 'KZ', 'TD', '证实', '可信']
    tm_list = ['TM', 'tm', '探明']
    kz_list = ['KZ', 'kz', '控制']
    td_list = ['TD', 'td', '推断']
    zs_list = ['证实']
    kx_list = ['可信']
    for root, dirs, files in os.walk(r"./word_new"):
        path_list = []
        for file in files:
            houzhui = file.split(".")[-1]
            # if houzhui=="doc":
            #     file = doc2Docx(root, file)
            path_list.append(os.path.join(root, file))
    # path = "./word/GQTCBG_420122402_乌龙泉矿生产矿山矿产资源国情调查报告.docx"
    for path in path_list:
        para_list = read_docx(path)
        if para_list:
            select_para(para_list)
    print(es_dataset)
    upload_data(es_dataset)
    jpype.shutdownJVM()  # 最后关闭JVM虚拟机
    # recognize("调整后的保有数据为：石灰岩TM:13477千吨，KZ:16015千吨，TD:2992千吨；白云岩TM:618千吨，KZ:7374千吨，TD:0千吨。调整变化量为：石灰岩保有TM:-871千吨，KZ:-360千吨；白云岩保有TM:-551千吨，KZ:-305千吨。")
